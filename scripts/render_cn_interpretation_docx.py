#!/usr/bin/env python
"""Render a Chinese paper interpretation Markdown file into a styled DOCX."""

from __future__ import annotations

import argparse
import re
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_COLOR_INDEX
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


CONTENT_WIDTH_CM = 16.2


def set_run_font(run, size_pt: float | None = None, bold: bool | None = None) -> None:
    font = run.font
    font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    if size_pt is not None:
        font.size = Pt(size_pt)
    if bold is not None:
        font.bold = bold


def set_paragraph_format(paragraph, first_line: bool = False) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(6)
    fmt.line_spacing = 1.25
    if first_line:
        fmt.first_line_indent = Cm(0.74)


def add_inline_markdown(paragraph, text: str, size: float = 11, bold_default: bool = False) -> None:
    parts = re.split(r"(\*\*[^*]+\*\*)", text)
    for part in parts:
        if not part:
            continue
        is_bold = part.startswith("**") and part.endswith("**")
        content = part[2:-2] if is_bold else part
        run = paragraph.add_run(content)
        set_run_font(run, size, bold_default or is_bold)


def set_cell_shading(paragraph, color: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color)
    p_pr.append(shd)


def set_cell_fill(cell, color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), color)


def set_cell_borders(cell, color: str = "111111", size: str = "18") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, margin_twips: str = "180") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for edge in ("top", "left", "bottom", "right"):
        element = margins.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            margins.append(element)
        element.set(qn("w:w"), margin_twips)
        element.set(qn("w:type"), "dxa")


def add_bottom_border(paragraph, color: str = "D8273A", size: str = "18") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "2")
    bottom.set(qn("w:color"), color)


def add_heading(document: Document, text: str, level: int) -> None:
    if level == 1:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(p)
        run = p.add_run(text)
        set_run_font(run, 18, True)
        run.font.color.rgb = RGBColor(31, 78, 121)
        return

    p = document.add_paragraph()
    set_paragraph_format(p)
    run = p.add_run(text)
    if level == 2:
        set_run_font(run, 14, True)
        run.font.color.rgb = RGBColor(31, 78, 121)
        set_cell_shading(p, "EAF2F8")
    else:
        set_run_font(run, 12, True)
        run.font.color.rgb = RGBColor(64, 64, 64)


def add_image(document: Document, markdown_path: Path, alt_text: str, image_ref: str) -> None:
    image_path = Path(image_ref)
    if not image_path.is_absolute():
        image_path = (markdown_path.parent / image_path).resolve()
    if not image_path.exists():
        p = document.add_paragraph()
        set_paragraph_format(p)
        run = p.add_run(f"[Image missing: {image_ref}]")
        set_run_font(run, 10, True)
        run.font.color.rgb = RGBColor(156, 0, 6)
        return

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(p)
    run = p.add_run()
    run.add_picture(str(image_path), width=Cm(CONTENT_WIDTH_CM))

    if alt_text:
        caption = document.add_paragraph()
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph_format(caption)
        caption_run = caption.add_run(alt_text)
        set_run_font(caption_run, 9, False)
        caption_run.font.color.rgb = RGBColor(89, 89, 89)


def is_internalization_card_heading(line: str) -> bool:
    return line.startswith("## ") and "文献内化卡片" in line


def add_internalization_card(document: Document, lines: list[str]) -> None:
    table = document.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    cell = table.cell(0, 0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    set_cell_fill(cell, "FFFDF3")
    set_cell_borders(cell)
    set_cell_margins(cell, "220")

    title = cell.paragraphs[0]
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(title)
    title_run = title.add_run("文献内化卡片")
    set_run_font(title_run, 18, True)
    title_run.font.name = "Microsoft YaHei"
    title_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    title_run.font.color.rgb = RGBColor(0, 0, 0)
    add_bottom_border(title)

    subtitle = cell.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_format(subtitle)
    subtitle_run = subtitle.add_run("一页读懂，一眼带走")
    set_run_font(subtitle_run, 9, False)
    subtitle_run.font.color.rgb = RGBColor(89, 89, 89)

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        p = cell.add_paragraph()
        set_paragraph_format(p)
        label_match = re.match(r"^([^：:]{1,18}[：:])(.+)$", line)
        if label_match:
            label, rest = label_match.groups()
            label_run = p.add_run(label)
            set_run_font(label_run, 11, True)
            label_run.font.name = "Microsoft YaHei"
            label_run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            label_run.font.color.rgb = RGBColor(0, 0, 0)
            label_run.font.highlight_color = WD_COLOR_INDEX.YELLOW

            add_inline_markdown(p, rest.strip(), 10.5)
        else:
            add_inline_markdown(p, line, 10.5)

    document.add_paragraph()


def render_markdown(input_path: Path, output_path: Path) -> None:
    document = Document()

    section = document.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)

    text = input_path.read_text(encoding="utf-8-sig")
    lines = text.splitlines()

    i = 0
    while i < len(lines):
        raw_line = lines[i]
        line = raw_line.strip().lstrip("\ufeff")
        if not line:
            i += 1
            continue

        if is_internalization_card_heading(line):
            card_lines: list[str] = []
            i += 1
            while i < len(lines):
                next_line = lines[i].strip().lstrip("\ufeff")
                if next_line.startswith("## "):
                    break
                card_lines.append(next_line)
                i += 1
            add_internalization_card(document, card_lines)
            continue

        image_match = re.match(r"^!\[([^\]]*)\]\(([^)]+)\)$", line)
        if image_match:
            add_image(document, input_path, image_match.group(1).strip(), image_match.group(2).strip())
        elif line.startswith("# "):
            add_heading(document, line[2:].strip(), 1)
        elif line.startswith("## "):
            add_heading(document, line[3:].strip(), 2)
        elif line.startswith("### "):
            add_heading(document, line[4:].strip(), 3)
        elif line.startswith("- "):
            p = document.add_paragraph(style="List Bullet")
            set_paragraph_format(p)
            add_inline_markdown(p, line[2:].strip(), 11)
        elif re.match(r"^\d+\.\s+", line):
            p = document.add_paragraph(style="List Number")
            set_paragraph_format(p)
            add_inline_markdown(p, re.sub(r"^\d+\.\s+", "", line), 11)
        else:
            p = document.add_paragraph()
            set_paragraph_format(p, first_line=True)
            label_match = re.match(r"^([^：:]{1,28}[：:])(.+)$", line)
            if label_match:
                label, rest = label_match.groups()
                label_run = p.add_run(label)
                set_run_font(label_run, 11, True)
                add_inline_markdown(p, rest.strip(), 11)
            else:
                add_inline_markdown(p, line, 11)

        i += 1

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("input_markdown", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    render_markdown(args.input_markdown, args.output_docx)


if __name__ == "__main__":
    main()
